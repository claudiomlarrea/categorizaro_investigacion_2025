# -*- coding: utf-8 -*-
import io, re, sys
from datetime import datetime
from typing import Dict, Tuple

import streamlit as st
import pandas as pd
import numpy as np

st.set_page_config(
    page_title="Valorador Automático de Currículum – Categorización de Investigadores",
    page_icon="📊",
    layout="wide"
)

# ──────────────────────────────────────────────────────────────────────────────
# UTILIDADES: importación perezosa para exportar Word
# ──────────────────────────────────────────────────────────────────────────────
Document = Pt = WD_ALIGN_PARAGRAPH = None
def _ensure_docx():
    global Document, Pt, WD_ALIGN_PARAGRAPH
    if Document is None:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    return Document, Pt, WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────────────────────────────────────
# CONSTANTES DE PUNTAJE / TOPES (mismos que definimos contigo)
# ──────────────────────────────────────────────────────────────────────────────
CATEGORIA_RANGOS = [
    ("I – Investigador Superior", 1000, 2000),
    ("II – Investigador Principal", 500, 999),
    ("III – Investigador Independiente", 300, 499),
    ("IV – Investigador Adjunto", 100, 299),
    ("V – Investigador Asistente", 1, 99),
    ("VI – Becario de Iniciación", 0, 0),
]
SECCIONES_MAX = {
    "Formación académica y complementaria": 450,
    "Cargos (docencia, gestión y otros)": 450,
    "Ciencia y Tecnología": 600,
    "Producciones y servicios": 600,
    "Otros antecedentes": 200,
}

# ──────────────────────────────────────────────────────────────────────────────
# EXTRACCIÓN DE TEXTO (PDF/Word)
# ──────────────────────────────────────────────────────────────────────────────
def read_pdf_text(file) -> str:
    from pypdf import PdfReader
    reader = PdfReader(file)
    out = []
    for p in reader.pages:
        try:
            out.append(p.extract_text() or "")
        except Exception:
            pass
    return "\n".join(out)

def read_docx_text(file) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file)
    return "\n".join(p.text for p in doc.paragraphs)

def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s.lower()).strip()

def safe_int(x, default=0):
    try:
        return int(x)
    except Exception:
        return default

# ──────────────────────────────────────────────────────────────────────────────
# HEURÍSTICAS (prototipo robusto y mejorable): cuentan ocurrencias por patrones
# ──────────────────────────────────────────────────────────────────────────────
def parse_counts_from_text(text: str) -> Tuple[Dict[str, int], Dict[str, str]]:
    """
    Devuelve:
      counts: valores 'estimados' para cada ítem
      notes:  pistas encontradas para auditoría
    Las heurísticas buscan palabras/expresiones frecuentes.
    """
    t = normalize(text)
    counts: Dict[str, int] = dict(
        # Formación
        doctorado=0, maestria=0, especializacion=0, diplomatura=0,
        posdoctorado=0, cursos_posgrado=0, idiomas=0, estancias=0, grados_extra=0,
        # Docencia / gestión
        doc_titular=0, doc_asoc=0, doc_adj=0, doc_jtp=0, doc_posgrad=0,
        gest_rector=0, gest_vice=0, gest_dec=0, gest_sec=0, gest_coord=0, oc=0,
        # CyT
        dir_inv=0, dir_tes=0, bec=0, dir_p=0, codir_p=0, part_p=0, coord_lin=0,
        tut=0, vinc=0, even=0, eg=0, ep=0, eprog=0, erev=0, einst=0, redes=0, prof=0,
        # Producciones/servicios
        art_ref=0, art_sin=0, libros=0, capitulos=0, doc_trab=0,
        pat_soft=0, procesos=0, serv_tec=0, informes=0,
        # Otros
        redes2=0, org_ev=0, gest_ed=0, prem_int=0, prem_nac=0, menc=0
    )
    notes: Dict[str, str] = {}

    # Formación
    counts["doctorado"]      = len(re.findall(r"\b(ph\.?d|doctor(a|ado) en)\b", t))
    counts["maestria"]       = len(re.findall(r"\b(maestr(ía|ia)|mag(í|i)ster)\b", t))
    counts["especializacion"]= len(re.findall(r"\bespecializaci(ó|o)n\b", t))
    counts["diplomatura"]    = len(re.findall(r"\bdiplomatur(a|as)\b", t))
    counts["posdoctorado"]   = len(re.findall(r"\bpos(doc|doctorado)\b", t))
    counts["cursos_posgrado"]= len(re.findall(r"(curso(s)? de (pos|post)grado|seminario de posgrado)", t))
    counts["idiomas"]        = len(re.findall(r"(toefl|ielts|cambridge|first certificate|dele|dalf|celu|b2|c1|c2|alicante)", t))
    counts["estancias"]      = len(re.findall(r"\b(estancia|pasant(í|i)a|sabbatical)\b", t))
    # si hay 2+ títulos de grado distintos (licenciado, ingeniero, médico...), suma 30 puntos directos
    multigrado = len(re.findall(r"\b(licenciado|ingeniero|abogado|m(é|e)dico|contador|arquitecto)\b", t))
    counts["grados_extra"]   = 30 if multigrado >= 2 else 0

    # Docencia (se cuentan cursos/roles mencionados; no inferimos años exactos)
    counts["doc_titular"] = len(re.findall(r"prof(\.|esor)? titular", t))
    counts["doc_asoc"]    = len(re.findall(r"prof(\.|esor)? asociado", t))
    counts["doc_adj"]     = len(re.findall(r"prof(\.|esor)? adjunto", t))
    counts["doc_jtp"]     = len(re.findall(r"(jtp|jefe de trabajos pr(á|a)cticos|ayudante)", t))
    counts["doc_posgrad"] = len(re.findall(r"(docencia|curso) de posgrado", t))

    # Gestión
    counts["gest_rector"] = 1 if re.search(r"\brector\b", t) else 0
    counts["gest_vice"]   = 1 if re.search(r"\b(vicerrector|directorio universitario)\b", t) else 0
    counts["gest_dec"]    = len(re.findall(r"\b(decano|director (de|del) (facultad|departamento|instituto))\b", t))
    counts["gest_sec"]    = len(re.findall(r"\bsecretar(i|í)a (acad(é|e)mica|investigaci(ó|o)n|extensi(ó|o)n)\b", t))
    counts["gest_coord"]  = len(re.findall(r"\bcoordinador(a)?|responsable\b", t))
    counts["oc"]          = len(re.findall(r"\b(comisi(ó|o)n|consejo|representante)\b", t))

    # CyT – formación RRHH
    counts["dir_inv"]  = len(re.findall(r"\b(direcci(ó|o)n|co-?direcci(ó|o)n) de (investigadores|becarios doctorales)\b", t))
    counts["dir_tes"]  = len(re.findall(r"\b(direcci(ó|o)n|co-?direcci(ó|o)n) de (tesis|tesistas)\b", t))
    counts["bec"]      = len(re.findall(r"\bbecario(s)?\b", t))

    # CyT – proyectos
    counts["dir_p"]    = len(re.findall(r"\b(dirigi(ó|o)|direcci(ó|o)n) (de )?proyecto(s)?\b", t))
    counts["codir_p"]  = len(re.findall(r"\bco-?direcci(ó|o)n (de )?proyecto(s)?\b", t))
    counts["part_p"]   = len(re.findall(r"\bparticipaci(ó|o)n en proyecto(s)?\b", t))
    counts["coord_lin"]= len(re.findall(r"\b(coordinaci(ó|o)n) (de )?l(í|i)nea(s)? (interdisciplinaria(s)?)\b", t))

    # Extensión
    counts["tut"]   = len(re.findall(r"\btutor(í|i)a(s)? (de )?(pasant(í|i)as|pr(á|a)cticas)\b", t))
    counts["vinc"]  = len(re.findall(r"\bvinculaci(ó|o)n|transferencia tecnol(ó|o)gica\b", t))
    counts["even"]  = len(re.findall(r"\b(congreso|jornada|simposio|encuentro)\b", t))  # se recorta luego por tope

    # Evaluación
    counts["eg"]    = len(re.findall(r"\btribunal (de )?tesis (de )?grado\b", t))
    counts["ep"]    = len(re.findall(r"\btribunal (de )?tesis (de )?posgrado\b", t))
    counts["eprog"] = len(re.findall(r"\bevaluaci(ó|o)n (de )?(programas|proyectos)\b", t))
    counts["erev"]  = len(re.findall(r"\b(reviewer|evaluaci(ó|o)n) (de )?(art(í|i)culos|revistas|congresos)\b", t))
    counts["einst"] = len(re.findall(r"\bevaluaci(ó|o)n institucional|organismo evaluador\b", t))

    # Otras CyT
    counts["redes"] = len(re.findall(r"\bred(es)? acad(é|e)micas|comit(é|e)s|sociedad cient(í|i)fica\b", t))
    counts["prof"]  = len(re.findall(r"\bejercicio profesional\b", t))

    # Producciones
    counts["art_ref"]  = len(re.findall(r"\b(art(í|i)culo|paper).*(scopus|wos|indexad|peer.?review|con referato)\b", t))
    counts["art_sin"]  = len(re.findall(r"\bart(í|i)culo\b", t)) - counts["art_ref"]
    counts["libros"]   = len(re.findall(r"\blibro(s)? (isbn)?\b", t))
    counts["capitulos"]= len(re.findall(r"\bcap(í|i)tulo(s)? de libro\b", t))
    counts["doc_trab"] = len(re.findall(r"\b(documento de trabajo|working paper|informe t(é|e)cnico)\b", t))

    counts["pat_soft"] = len(re.findall(r"\b(patente|modelo de utilidad|software registrado)\b", t))
    counts["procesos"] = len(re.findall(r"\b(proceso|innovaci(ó|o)n)\b", t))
    counts["serv_tec"] = len(re.findall(r"\bservicio(s)? tecn(ó|o)logico(s)?\b", t))
    counts["informes"] = len(re.findall(r"\binforme(s)? t(é|e)cnico(s)?\b", t))

    # Otros antecedentes
    counts["redes2"]   = len(re.findall(r"\bred(es)? (acad(é|e)micas|cient(í|i)ficas)\b", t))
    counts["org_ev"]   = len(re.findall(r"\b(organizador|comit(é|e) organizador)\b", t))
    counts["gest_ed"]  = len(re.findall(r"\b(editor|comit(é|e) editorial|gesti(ó|o)n editorial)\b", t))
    counts["prem_int"] = len(re.findall(r"\bpremio(s)? (internacional(es)?)\b", t))
    counts["prem_nac"] = len(re.findall(r"\bpremio(s)? (nacional(es)?)\b", t))
    counts["menc"]     = len(re.findall(r"\bmenci(ó|o)n(es)? honor(í|i)fica(s)?\b", t))

    # Notas de auditoría simples
    notes["texto_len"] = f"{len(t)} caracteres"
    return counts, notes

# ──────────────────────────────────────────────────────────────────────────────
# CÁLCULO DE PUNTAJES (usa counts detectados + topes institucionales)
# ──────────────────────────────────────────────────────────────────────────────
def compute_scores(c: Dict[str, int]) -> Dict[str, int]:
    # 1) Formación (máx. 450)
    form_total = 0
    form_total += min(c["doctorado"] * 250, 375)
    form_total += min(c["maestria"] * 150, 225)
    form_total += min(c["especializacion"] * 70, 105)
    form_total += min(c["diplomatura"] * 50, 100)
    form_total += min(c["posdoctorado"] * 100, 100)
    form_total += min(c["cursos_posgrado"] * 5, 75)
    form_total += min(c["idiomas"] * 10, 50)
    form_total += min(c["estancias"] * 20, 60)
    form_total += min(c["grados_extra"], 30)
    form_total = min(int(form_total), SECCIONES_MAX["Formación académica y complementaria"])

    # 2) Cargos (máx. 500)
    docencia = min(c["doc_titular"] * 30, 150) + min(c["doc_asoc"] * 25, 125) + \
               min(c["doc_adj"] * 20, 100) + min(c["doc_jtp"] * 10, 50) + \
               min(c["doc_posgrad"] * 20, 100)
    docencia = min(docencia, 300)

    gestion = (100 if c["gest_rector"] else 0) + (80 if c["gest_vice"] else 0) + \
              min(c["gest_dec"], 60) + min(c["gest_sec"], 60) + min(c["gest_coord"], 40)
    gestion = min(gestion, 200)

    otros_cargos = min(c["oc"] * 10, 50)
    cargos = min(docencia + gestion + otros_cargos, SECCIONES_MAX["Cargos (docencia, gestión y otros)"])

    # 3) CyT (máx. 500)
    form_cyt = min(c["dir_inv"], 90) + min(c["dir_tes"], 50) + min(c["bec"] * 20, 40)
    form_cyt = min(form_cyt, 150)

    proyectos = min(c["dir_p"] * 50, 150) + min(c["codir_p"] * 30, 90) + \
                min(c["part_p"] * 20, 60) + min(c["coord_lin"] * 20, 20)
    proyectos = min(proyectos, 150)

    extension = min(c["tut"] * 10, 20) + min(c["vinc"] * 15, 45) + min(c["even"], 100)
    extension = min(extension, 60)

    evaluacion = min(c["eg"] * 5, 20) + min(c["ep"] * 10, 30) + min(c["eprog"] * 10, 30) + \
                 min(c["erev"] * 10, 30) + min(c["einst"] * 10, 30)
    evaluacion = min(evaluacion, 100)

    otras_cyt = min(c["redes"] * 20, 60) + min(c["prof"] * 5, 20)
    otras_cyt = min(otras_cyt, 60)

    cyt = min(form_cyt + proyectos + extension + evaluacion + otras_cyt, SECCIONES_MAX["Ciencia y Tecnología"])

    # 4) Producciones y servicios (máx. 350)
    publicaciones = min(c["art_ref"] * 20, 140) + min(c["art_sin"] * 10, 80) + \
                    min(c["libros"] * 40, 80) + min(c["capitulos"] * 20, 60) + \
                    min(c["doc_trab"] * 10, 30)
    publicaciones = min(publicaciones, 300)

    desarrollos = min(c["pat_soft"] * 30, 60) + min(c["procesos"] * 20, 60)
    desarrollos = min(desarrollos, 100)

    servicios = min(c["serv_tec"] * 20, 40) + min(c["informes"] * 10, 20)
    servicios = min(servicios, 40)

    prod = min(publicaciones + desarrollos + servicios, SECCIONES_MAX["Producciones y servicios"])

    # 5) Otros antecedentes (máx. 200)
    redes_gestion = min(c["redes2"] * 10, 30) + min(c["org_ev"] * 20, 60) + min(c["gest_ed"], 60)
    redes_gestion = min(redes_gestion, 150)

    premios = min(c["prem_int"] * 50, 100) + min(c["prem_nac"] * 20, 100) + min(c["menc"] * 20, 100)
    premios = min(premios, 100)

    otros = min(redes_gestion + premios, SECCIONES_MAX["Otros antecedentes"])

    total = form_total + cargos + cyt + prod + otros
    return dict(
        formacion=form_total, cargos=cargos, cyt=cyt, prod=prod, otros=otros, total=total
    )

def determinar_categoria(total: int) -> str:
    for nombre, lo, hi in CATEGORIA_RANGOS:
        if lo <= total <= hi or (nombre == "I – Investigador Superior" and total > hi):
            return nombre
    return "VI – Becario de Iniciación"

def tag_categoria(cat: str) -> str:
    base = {"I": ("#065f46", "#ecfdf5"), "II": ("#064e3b", "#ecfdf5"),
            "III": ("#1e40af", "#eff6ff"), "IV": ("#7c2d12", "#fffbeb"),
            "V": ("#7f1d1d", "#fef2f2"), "VI": ("#334155", "#f1f5f9")}
    key = cat.split(" ")[0]
    fg, bg = base.get(key, ("#334155", "#f1f5f9"))
    return f"<span style='background:{bg}; color:{fg}; padding:6px 10px; border-radius:999px; font-weight:600;'>{cat}</span>"

# ──────────────────────────────────────────────────────────────────────────────
# UI
# ──────────────────────────────────────────────────────────────────────────────
st.markdown(
    """
    <div style="padding: 18px; border-radius: 14px; background: #0f172a; color: white; text-align:center;">
      <h1 style="margin: 0 0 6px 0;">Universidad Católica de Cuyo</h1>
      <div style="opacity:.9;">Secretaría de Investigación · <b>Valorador Automático de Currículum</b></div>
    </div>
    """,
    unsafe_allow_html=True
)
st.write("")

colA, colB = st.columns([2,1])
with colA:
    uploaded = st.file_uploader("Subí el CV del postulante (PDF o Word .docx)", type=["pdf", "docx"])
with colB:
    nombre = st.text_input("Nombre y Apellido", "")
    unidad = st.text_input("Unidad Académica / Instituto", "")
    fecha_eval = st.date_input("Fecha de evaluación", datetime.now())

if not uploaded:
    st.info("Cargá un PDF/Word del CV para valorar automáticamente.")
    st.stop()

# ── EXTRACCIÓN DE TEXTO ──────────────────────────────────────────────────────
try:
    if uploaded.name.lower().endswith(".pdf"):
        raw_text = read_pdf_text(uploaded)
    else:
        raw_text = read_docx_text(uploaded)
except Exception as e:
    st.error("No se pudo leer el archivo. ¿Es un PDF escaneado? (no se admite OCR en esta versión).")
    st.exception(e)
    st.stop()

if not raw_text.strip():
    st.warning("No se extrajo texto. Si el PDF es una imagen/escaneo, requeriría OCR.")
    st.stop()

with st.expander("Vista previa del texto extraído", expanded=False):
    st.text_area("Texto", raw_text[:20000], height=220)

# ── PARSEO + CÁLCULO ─────────────────────────────────────────────────────────
counts, notes = parse_counts_from_text(raw_text)
scores = compute_scores(counts)
categoria = determinar_categoria(scores["total"])

st.success("Valoración completada automáticamente ✅")

# ── RESUMEN ──────────────────────────────────────────────────────────────────
st.markdown("---")
c1, c2, c3, c4 = st.columns(4)
c1.metric("Formación", scores["formacion"])
c2.metric("Cargos", scores["cargos"])
c3.metric("CyT", scores["cyt"])
c4.metric("Producciones/Servicios", scores["prod"])
c5, c6 = st.columns([1,2])
c5.metric("Otros antecedentes", scores["otros"])
c6.metric("TOTAL", scores["total"])

st.markdown(f"**Categoría alcanzada:** {tag_categoria(categoria)}", unsafe_allow_html=True)

# Tabla de secciones
resumen = pd.DataFrame([
    {"Sección": "Formación académica y complementaria", "Puntaje": scores["formacion"], "Tope": SECCIONES_MAX["Formación académica y complementaria"]},
    {"Sección": "Cargos (docencia, gestión y otros)", "Puntaje": scores["cargos"], "Tope": SECCIONES_MAX["Cargos (docencia, gestión y otros)"]},
    {"Sección": "Ciencia y Tecnología", "Puntaje": scores["cyt"], "Tope": SECCIONES_MAX["Ciencia y Tecnología"]},
    {"Sección": "Producciones y servicios", "Puntaje": scores["prod"], "Tope": SECCIONES_MAX["Producciones y servicios"]},
    {"Sección": "Otros antecedentes", "Puntaje": scores["otros"], "Tope": SECCIONES_MAX["Otros antecedentes"]},
    {"Sección": "TOTAL", "Puntaje": scores["total"], "Tope": 2000},
])
st.dataframe(resumen, use_container_width=True)

with st.expander("Detalle de detectores (auditoría)", expanded=False):
    df_counts = pd.DataFrame(sorted(counts.items()), columns=["Ítem", "Conteo detectado"])
    st.dataframe(df_counts, use_container_width=True)
    st.caption(f"Notas: {notes.get('texto_len','')}")

# ── EXPORTABLES ──────────────────────────────────────────────────────────────
colx1, colx2 = st.columns(2)
with colx1:
    # Excel
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
        resumen.to_excel(writer, index=False, sheet_name="Resultados")
        pd.DataFrame(sorted(counts.items()), columns=["Ítem", "Conteo"]).to_excel(writer, index=False, sheet_name="Detectores")
    st.download_button(
        "⬇️ Descargar Excel (resultados + detectores)",
        out.getvalue(),
        file_name=f"Valorador_Investigadores_Auto_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

with colx2:
    # Word (import perezosa)
    Document, Pt, WD_ALIGN_PARAGRAPH = _ensure_docx()
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(11)
    h = doc.add_paragraph("Universidad Católica de Cuyo – Secretaría de Investigación")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.runs[0].bold = True
    doc.add_paragraph("Valorador Automático de Currículum – Categorización de Investigadores")
    info = doc.add_paragraph(f"Postulante: {nombre if nombre else '-'} | Unidad: {unidad if unidad else '-'} | Fecha: {fecha_eval.strftime('%Y-%m-%d')}")
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Categoría alcanzada: {categoria}")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells; hdr[0].text = "Sección"; hdr[1].text = "Puntaje"; hdr[2].text = "Tope"
    for _, row in resumen.iterrows():
        r = table.add_row().cells
        r[0].text = str(row["Sección"]); r[1].text = str(int(row["Puntaje"])); r[2].text = str(int(row["Tope"]))
    doc.add_paragraph("")
    doc.add_paragraph("Auditoría (detectores):")
    for k, v in sorted(counts.items()):
        doc.add_paragraph(f"• {k}: {v}")
    wb = io.BytesIO(); doc.save(wb)
    st.download_button(
        "⬇️ Descargar Informe Word",
        wb.getvalue(),
        file_name=f"Valorador_Investigadores_Auto_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
